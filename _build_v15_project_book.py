from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "_generated_assets_v15"
OUT = ROOT / "Driver_Moodle_FINAL_v15_REBUILT.docx"
QA_OUT = ROOT / "Driver_Moodle_EXAM_QA_v15.docx"

BLUE = "1F4E79"
TEAL = "2D8C9E"
GREEN = "2F7D5A"
ORANGE = "B8741A"
GRAY = "5B6573"
LIGHT_BLUE = "EAF3F8"
LIGHT_GRAY = "F5F7F9"
WHITE = "FFFFFF"

bookmark_id = 1200
pending_page_break = False


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = int(sum(widths_inches) * 1440)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = int(widths_inches[idx] * 1440)
            cell.width = Inches(widths_inches[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=11, bold=None, color=None, name="Arial", rtl=False):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if rtl:
        rtl_node = r_pr.find(qn("w:rtl"))
        if rtl_node is None:
            rtl_node = OxmlElement("w:rtl")
            r_pr.append(rtl_node)


def set_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_ltr(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is not None:
        p_pr.remove(bidi)


def add_para(doc, text="", size=11, bold=False, color=None, after=6,
             before=0, align=WD_ALIGN_PARAGRAPH.RIGHT, rtl=True, style=None):
    p = doc.add_paragraph(style=style)
    if rtl:
        set_rtl(p, align)
    else:
        set_ltr(p, align)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color, rtl=rtl)
    return p


def add_heading(doc, text, level=1, bookmark=None):
    global pending_page_break
    p = doc.add_heading(text, level=level)
    set_rtl(p)
    if pending_page_break:
        p.paragraph_format.page_break_before = True
        pending_page_break = False
    p.paragraph_format.keep_with_next = True
    for r in p.runs:
        set_run_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11),
                     bold=True, color=BLUE if level <= 2 else TEAL, rtl=True)
    if bookmark:
        add_bookmark(p, bookmark)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_rtl(p)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, rtl=True)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_rtl(p)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, rtl=True)
    return p


def add_bookmark(paragraph, name):
    global bookmark_id
    bookmark_id += 1
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_link(paragraph, text, anchor):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_break(doc):
    # Apply the break to the next heading. This avoids blank PDF pages when the
    # previous element is a table or an image close to the bottom margin.
    global pending_page_break
    pending_page_break = True


def add_picture(doc, path, width=6.05, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        cap = add_para(doc, caption, size=9.5, color=GRAY, after=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        return cap
    return p


def add_code(doc, code, caption=None):
    if caption:
        add_para(doc, caption, size=10, bold=True, color=BLUE, after=3)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_fixed(table, [6.15])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F5F7")
    p = cell.paragraphs[0]
    set_ltr(p)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(code.strip())
    set_run_font(r, size=8.2, name="Consolas")
    add_para(doc, "", after=4)
    return table


def add_table(doc, headers, rows, widths, font_size=9.5, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_fixed(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(h)
        set_run_font(r, size=font_size, bold=True, color=BLUE, rtl=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            set_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
            r = p.add_run(str(val))
            set_run_font(r, size=font_size, rtl=True)
    set_table_fixed(table, widths)
    add_para(doc, "", after=4)
    return table


def add_requirements_table(doc):
    add_heading(doc, "בדיקת דרישות המחוון", 1, "requirements")
    add_para(doc,
             "כדי שיהיה קל לבדוק את התיק, חילקתי בין דרישות החובה לבין דרישות הבחירה. "
             "בכל שורה יש קישור שמוביל ישירות להסבר המתאים בתוך המסמך.",
             after=8)
    rows = [
        ("חובה", "1", "מערכת לקוח-שרת", "שלושה לקוחות פונים לאותו שירות WCF", "server"),
        ("חובה", "2", "בסיס נתונים יחסי", "מסד Access עם טבלאות וקשרים", "database"),
        ("חובה", "3", "שכבות בפרויקט", "Model, ViewDB, BusinessLogic ושירות", "architecture"),
        ("חובה", "4", "פעולות CRUD", "שליפה, הוספה, עדכון ומחיקה דרך ViewDB", "crud"),
        ("חובה", "5", "לקוח WPF", "יישום למחשב הביתי", "wpf"),
        ("חובה", "6", "לקוח Web", "דפי Razor Pages בדפדפן", "web"),
        ("חובה", "7", "לקוח MAUI", "יישום Android לפעולות קצרות מהטלפון", "maui"),
        ("חובה", "8", "Use Case והרשאות", "תלמיד, מורה ומנהל", "usecase"),
        ("בחירה", "9", "ירושה ומחלקות גנריות", "Base, BaseDB ורשימות אובייקטים", "model"),
        ("בחירה", "9", "ניהול מספר משתמשים", "הפרדה לפי מזהה משתמש וסוג משתמש", "notifications"),
        ("בחירה", "10", "IValueConverter ו-XAML", "המרת מידע לתצוגה בלקוח WPF", "converter"),
        ("בחירה", "10", "Validation", "בדיקות קלט לפני שמירה", "security"),
        ("בחירה", "10", "שירות חיצוני", "WCF מחבר בין הלקוחות לשרת", "wcf"),
        ("בחירה", "10", "SQL Injection", "OleDbParameter במקום חיבור טקסט", "security"),
        ("בחירה", "10", "Async / Tasks", "קריאות אסינכרוניות בלקוח MAUI", "async"),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["סוג", "מספר", "דרישה", "איפה מימשתי", "מעבר לפירוט"]
    widths = [0.72, 0.52, 1.52, 2.82, 0.70]
    set_table_fixed(table, widths)
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "DCEEEF")
        p = cell.paragraphs[0]
        set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(text)
        set_run_font(r, size=8.5, bold=True, color=BLUE, rtl=True)
    for kind, number, req, where, anchor in rows:
        cells = table.add_row().cells
        values = [kind, number, "✓ " + req, where]
        for i, text in enumerate(values):
            p = cells[i].paragraphs[0]
            set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1) else WD_ALIGN_PARAGRAPH.RIGHT)
            r = p.add_run(text)
            set_run_font(r, size=8.2, rtl=True)
        p = cells[4].paragraphs[0]
        set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        add_internal_link(p, "לפירוט", anchor)
    set_table_fixed(table, widths)


def add_toc(doc):
    add_heading(doc, "תוכן עניינים", 1)
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "יש לעדכן את תוכן העניינים ב-Word."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, txt, fld_end])


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("אלי רפופורט | Driver Moodle | עמוד ")
    set_run_font(r, size=8.5, color=GRAY, rtl=True)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.extend([fld_begin, instr, fld_end])


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = True
    add_footer(section)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 12, TEAL, 8, 4),
    ):
        st = styles[style_name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        st._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def cover(doc):
    add_para(doc, "מקיף ז' - הקריה, אשדוד", size=12, bold=True, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_picture(doc, ASSETS / "makifz_logo.png", width=1.0)
    add_para(doc, "מגמת הנדסת תוכנה 5 יח\"ל", size=12, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, "Driver Moodle", size=28, bold=True, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=3, rtl=False)
    add_para(doc, "מערכת לניהול שיעורי נהיגה", size=18, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "פרויקט גמר - שירותי רשת, בסיס נתונים ותכנות אסינכרוני",
             size=11.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_picture(doc, ASSETS / "driver_moodle_cover_visual.png", width=3.75)
    add_para(doc, "שם התלמיד: אלי רפופורט", size=11.5, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "תעודת זהות: 327588232", size=11.5,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "שם המנחה: אולגה גרדלמן", size=11.5,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "תאריך הגשה: יוני 2026", size=11.5,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=0)


def build_main():
    doc = Document()
    configure_doc(doc)
    cover(doc)
    add_page_break(doc)
    add_toc(doc)
    add_requirements_table(doc)

    add_page_break(doc)
    add_heading(doc, "פרק 1: מבוא ותרחישי שימוש", 1, "intro")
    add_heading(doc, "למה בחרתי בפרויקט", 2)
    add_para(doc,
             "בחרתי לבנות את Driver Moodle בגלל החוויה האישית שלי בזמן שלמדתי נהיגה. "
             "לא תמיד היה לי קל לקבוע שיעור. לפעמים קבענו דרך הודעות, לפעמים שעה השתנתה, "
             "ולא תמיד היה ברור כמה שיעורים כבר עשיתי ומה נשאר לשלם. רציתי לרכז את הדברים במקום אחד.")
    add_para(doc,
             "המערכת היא מערכת לקוח-שרת לניהול שיעורי נהיגה. הנתונים לא נשמרים בתוך מסך מסוים. "
             "כל לקוח שולח בקשות לשירות הרשת, ורק השרת עובד מול בסיס הנתונים.")
    add_heading(doc, "מי משתמש במערכת", 2)
    add_bullet(doc, "תלמיד: קובע שיעורים, צופה בשיעורים, רואה תשלומים, מקבל התראות ושולח פנייה.")
    add_bullet(doc, "מורה: מנהל לוח זמנים, רואה תלמידים, מאשר תשלומים, שולח התראות ורואה דוחות.")
    add_bullet(doc, "מנהל: מנהל משתמשים ומטפל בפניות תמיכה.")
    add_heading(doc, "Use Case", 2, "usecase")
    add_para(doc,
             "חילקתי את הפעולות לפי סוג המשתמש. כך כל משתמש רואה את המסכים שמתאימים לתפקיד שלו, "
             "ולא מקבל גישה למסכים שלא שייכים אליו.")
    add_picture(doc, ASSETS / "usecase_roles_portrait.png", width=5.45,
                caption="תרשים 1: הפעולות המרכזיות לפי סוג המשתמש.")

    add_page_break(doc)
    add_heading(doc, "פרק 2: בסיס הנתונים", 1, "database")
    add_para(doc,
             "בסיס הנתונים נשמר בקובץ UsersDataBase.accdb מסוג Access. בחרתי ב-Access כי הוא מתאים "
             "להיקף של פרויקט לימודי, מאפשר לעבוד עם טבלאות וקשרים, וקל להציג אותו בזמן הבחינה.")
    add_para(doc,
             "לא שמרתי את כל המידע בטבלה אחת. הפרדתי בין תלמידים, מורים, שיעורים, תשלומים, "
             "התראות ופניות תמיכה. כך אפשר לעדכן נתון אחד בלי לפגוע בנתונים אחרים.")
    add_heading(doc, "טבלאות מרכזיות", 2)
    add_table(doc, ["טבלה", "מה נשמר בה"], [
        ("Teacher", "פרטי מורה, מחיר שיעור, הרשאת מנהל ואמצעי תשלום."),
        ("Student", "פרטי תלמיד, המורה שלו, מצב אישור ומחיר מותאם."),
        ("Lessons", "שיעור, תלמיד, מורה, תאריך, שעה ומצב תשלום."),
        ("Payments", "תשלום, סכום, תלמיד, שיעור ומצב התשלום."),
        ("Notifications", "שולח, מקבל, כותרת, הודעה ומצב קריאה."),
        ("SupportTickets", "פניית תמיכה, נושא, תיאור, מצב ועדיפות."),
        ("TicketMessages", "הודעות שנכתבו בתוך פנייה."),
    ], [1.45, 4.75], font_size=9.5)
    add_heading(doc, "קשרים בין הטבלאות", 2, "relationships")
    add_para(doc,
             "בתרשים הבא מופיעים הקשרים המרכזיים שנשמרים במסד. לדוגמה, תלמיד מקושר למורה, "
             "שיעור מקושר לתלמיד ולמורה, ותשלום מקושר לתלמיד ולשיעור.")
    add_picture(doc, ASSETS / "db_relationships_core_portrait.png", width=5.75,
                caption="תרשים 2: קשרים מרכזיים במסד הנתונים. הקווים הרציפים מייצגים קשרים שמוגדרים במסד.")
    add_heading(doc, "קשרים לוגיים נוספים", 2)
    add_para(doc,
             "יש גם טבלאות שבהן הקישור נשמר באמצעות מזהים ונבדק בקוד השירות. לדוגמה, התראות "
             "נשמרות לפי מזהה משתמש וסוג משתמש. בתרשים הבא סימנתי קשרים כאלה בקו מקווקו כדי לא להציג אותם "
             "כאילו הם קשרי מפתח זר שהוגדרו במסד.")
    add_picture(doc, ASSETS / "db_relationships_logical_portrait.png", width=5.75,
                caption="תרשים 3: טבלאות וקשרים לוגיים שמנוהלים בקוד השירות.")
    add_heading(doc, "גישה לבסיס הנתונים", 2, "crud")
    add_para(doc,
             "הלקוחות לא ניגשים ישירות ל-Access. העבודה עוברת דרך שכבת ViewDB. המחלקה BaseDB "
             "מרכזת את החיבור, את מנגנון השליפה ואת מנגנון השמירה. המחלקות האחרות משתמשות בה.")
    add_code(doc, """
protected virtual List<Base> Select(string sqlCommandTxt,
                                    params OleDbParameter[] parameters)
{
    connection.Open();
    command.CommandText = sqlCommandTxt;
    command.Parameters.Clear();
    command.Parameters.AddRange(parameters);
    reader = command.ExecuteReader();
    ...
}

protected int SaveChanges(string commandText,
                          params OleDbParameter[] parameters)
{
    cmd.CommandText = commandText;
    cmd.Parameters.AddRange(parameters);
    connection.Open();
    return cmd.ExecuteNonQuery();
}""", "קטע קוד: מנגנוני Select ו-SaveChanges במחלקה BaseDB.")
    add_para(doc,
             "באמצעות שני המנגנונים האלה מתבצעות פעולות CRUD: שליפה, הוספה, עדכון ומחיקה. "
             "הפקודה משתנה לפי הפעולה, אבל פתיחת החיבור והטיפול בפרמטרים נשארים במקום אחד.")

    add_page_break(doc)
    add_heading(doc, "פרק 3: צד השרת", 1, "server")
    add_heading(doc, "מבנה כללי", 2, "architecture")
    add_para(doc,
             "צד השרת מחולק לפרויקטים נפרדים. לכל פרויקט יש תפקיד ברור. החלוקה עוזרת לשמור "
             "על סדר ומונעת מצב שבו מסך הלקוח מכיל גם קוד של בסיס הנתונים.")
    add_picture(doc, ASSETS / "architecture_portrait.png", width=5.65,
                caption="תרשים 4: מבנה המערכת. כל הלקוחות עובדים מול אותו שירות.")
    add_picture(doc, ASSETS / "solution_structure_portrait.png", width=5.65,
                caption="תרשים 5: מבנה הפתרון לפי התיקיות והפרויקטים האמיתיים בפרויקט.")
    add_heading(doc, "פרויקט Model", 2, "model")
    add_para(doc,
             "בפרויקט Model נמצאות מחלקות שמייצגות מידע שעובר בין הלקוח לשרת ובחזרה. "
             "רוב המחלקות מכילות Properties. לדוגמה, UserInfo מייצגת משתמש, Notification מייצגת התראה, "
             "ו-Course מייצגת קורס.")
    add_para(doc,
             "מחלקות מרכזיות יורשות מהמחלקה Base. במחלקה Base נמצא Id משותף. כך לכל אובייקט יש מזהה "
             "שאפשר לקשר לרשומה המתאימה במסד הנתונים.")
    add_picture(doc, ASSETS / "uml_classes_portrait.png", width=5.75,
                caption="תרשים 6: UML של המחלקות המרכזיות. החצים מצביעים למחלקת האב או לממשק.")
    add_heading(doc, "פרויקט ViewDB", 2, "viewdb")
    add_para(doc,
             "בפרויקט ViewDB נמצאות המחלקות שעובדות מול בסיס הנתונים. BaseDB אחראית על פתיחת החיבור "
             "ועל פעולות בסיסיות. מחלקות כמו UserDB, LessonsDB, PaymentDB ו-NotificationDB אחראיות "
             "כל אחת על טבלה או נושא מסוים.")
    add_code(doc, """
public List<Lessons> GetAllStudentLessons(int sid)
{
    string sql = "SELECT * FROM [Lessons] WHERE StudentID = ?";
    return Select(sql, new OleDbParameter("@sid", sid))
        .OfType<Lessons>()
        .ToList();
}

public void CancelLesson(int lessonId)
{
    string sql = "UPDATE [Lessons] SET Canceled = 1 WHERE LessonId = ?";
    SaveChanges(sql, new OleDbParameter("@lessonId", lessonId));
}""", "קטע קוד: דוגמה לשליפה ולעדכון דרך LessonsDB.")
    add_heading(doc, "פרויקט BusinessLogic", 2, "businesslogic")
    add_para(doc,
             "בפרויקט BusinessLogic נמצאות מחלקות עזר שמרכזות עבודה על רשימות ונתונים. "
             "לדוגמה, קיימות מחלקות עבור רשימת משתמשים, רשימת תלמידים ורשימת ערים. "
             "השכבה מיועדת ללוגיקה שלא שייכת למסך ולא לשאילתת SQL אחת.")
    add_heading(doc, "שירות WCF", 2, "wcf")
    add_para(doc,
             "בחרתי בשירות WCF כדי שכל הלקוחות יעבדו מול אותו שרת. IService1 הוא הממשק שמגדיר "
             "אילו פעולות הלקוח יכול לבקש. Service1 מממש את הפעולות ומפנה אותן למחלקה המתאימה ב-ViewDB.")
    add_code(doc, """
[OperationContract]
void AddLessonForStudent(int sid, string Date, string time);

[OperationContract]
List<Lessons> GetAllStudentLessons(int id);

[OperationContract]
List<Notification> GetUserNotifications(int userId, string userType);""",
             "קטע קוד: דוגמאות לפעולות שהוגדרו בממשק IService1.")
    add_heading(doc, "קובצי App.config והגדרות הרשת", 2, "appconfig")
    add_para(doc,
             "כתובת השירות לא כתובה בכל מסך בנפרד. היא נשמרת בקובצי App.config. "
             "בצד הלקוח מוגדרת הכתובת שאליה מתחברים, ובצד השרת מוגדרת כתובת הבסיס של השירות. "
             "אם כתובת השרת משתנה, מעדכנים אותה בקובץ ההגדרות.")
    add_code(doc, """
<endpoint
  address="http://192.168.1.136:8733/Design_Time_Addresses/WcfServiceLibrary1/Service1/"
  binding="basicHttpBinding"
  contract="driver.IService1" />

<add baseAddress="http://192.168.1.136:8733/Design_Time_Addresses/WcfServiceLibrary1/Service1/" />""",
             "קטע קוד: כתובת השירות בקובץ App.config.")

    add_page_break(doc)
    add_heading(doc, "פרק 4: צד הלקוח", 1, "clients")
    add_para(doc,
             "למערכת יש שלושה לקוחות. הם לא נועדו להחליף אחד את השני. כל לקוח מתאים למצב שימוש אחר, "
             "אבל כולם פונים לאותו שירות רשת ועובדים על אותם נתונים.")
    add_table(doc, ["לקוח", "מי משתמש בו", "למה בחרתי בו"], [
        ("WPF", "תלמיד, מורה ומנהל מהמחשב הביתי", "נוח לעבודה מלאה עם מסכים גדולים והרבה פעולות."),
        ("Web", "תלמיד, מורה או מנהל שנכנסים מדפדפן", "לא דורש התקנה ומתאים לכניסה מהירה."),
        ("MAUI / Android", "בעיקר תלמיד ומורה מהטלפון", "מתאים לפעולות קצרות כמו צפייה, תשלום והתראות."),
    ], [1.25, 2.25, 2.70], font_size=9.5)
    add_heading(doc, "לקוח WPF", 2, "wpf")
    add_para(doc,
             "יישום WPF הוא הלקוח המרכזי למחשב. בחרתי בו כדי שתלמידים ומורים יוכלו לעבוד בנוחות "
             "מהבית. תלמיד יכול לקבוע שיעור, לראות תשלומים והתראות. מורה יכול לנהל תלמידים, לוח זמנים ודוחות.")
    add_picture(doc, ROOT / "nav-map" / "screenshots" / "StudentUI.png", width=6.15,
                caption="תמונה 1: מסך הבית של תלמיד ביישום WPF.")
    add_picture(doc, ROOT / "nav-map" / "screenshots" / "TeacherUI.png", width=6.15,
                caption="תמונה 2: מסך הבית של מורה ביישום WPF.")
    add_heading(doc, "Binding, XAML ו-IValueConverter", 3, "converter")
    add_para(doc,
             "מסכי WPF נכתבו ב-XAML. בחלק מהמסכים השתמשתי ב-Binding כדי לחבר בין נתונים לפקדים. "
             "בנוסף קיימת המחלקה ImgConventer שמממשת IValueConverter וממירה מידע לצורה שאפשר להציג במסך.")
    add_code(doc, """
public class ImgConventer : IValueConverter
{
    public object Convert(object value, Type targetType,
                          object parameter, CultureInfo culture)
    {
        ...
    }
}""", "קטע קוד: מחלקת הממיר בלקוח WPF.")
    add_heading(doc, "לקוח Web", 2, "web")
    add_para(doc,
             "לקוח Web נכתב עם Razor Pages. בחרתי בו כדי לאפשר כניסה מדפדפן בלי התקנה. "
             "יש בו דפים נפרדים לתלמיד, למורה ולמנהל, לדוגמה ScheduleLesson, Payments, Notifications "
             "ו-ManageUsers.")
    add_heading(doc, "לקוח MAUI / Android", 2, "maui")
    add_para(doc,
             "לקוח MAUI מיועד בעיקר לטלפון Android. בחרתי בו כדי לבצע פעולות קצרות גם כשלא נמצאים ליד מחשב. "
             "תלמיד יכול לצפות בשיעורים ולקבוע שיעור. מורה יכול לראות תשלומים והתראות.")
    add_heading(doc, "Async / Tasks", 3, "async")
    add_para(doc,
             "בטלפון חשוב שהמסך לא ייתקע בזמן פנייה לשרת. לכן בלקוח MAUI השתמשתי בפעולות אסינכרוניות "
             "וב-await. המשתמש יכול להמשיך לקבל תגובה מהמסך בזמן שהבקשה נשלחת.")
    add_code(doc, """
var lessons = await ServiceHelper.CallAsync(
    srv => srv.GetAllTeacherLessonsAsync(AppState.UserId));

await ServiceHelper.CallAsync(
    srv => srv.MarkLessonPaidAsync(lessonId));""",
             "קטע קוד: קריאה אסינכרונית לשירות מתוך לקוח MAUI.")

    add_page_break(doc)
    add_heading(doc, "פרק 5: תהליכים מרכזיים ואבטחה", 1)
    add_heading(doc, "קביעת שיעור", 2, "schedule_flow")
    add_para(doc,
             "כאשר תלמיד קובע שיעור, הלקוח שולח לשירות את מזהה התלמיד, התאריך והשעה. "
             "Service1 מעביר את הפעולה ל-LessonsDB. שם נבדק המורה של התלמיד ונשמרת רשומה חדשה בטבלת Lessons.")
    add_picture(doc, ASSETS / "sequence_schedule_lesson_portrait.png", width=5.75,
                caption="תרשים 7: רצף הפעולות כאשר תלמיד קובע שיעור.")
    add_heading(doc, "התראות", 2, "notifications")
    add_para(doc,
             "ההתראות נשמרות בטבלת Notifications. לכל התראה יש מזהה שולח, מזהה מקבל וסוג משתמש. "
             "כך השירות יודע להחזיר לכל משתמש רק את ההתראות שמתאימות לו. קיימות גם פעולות לסימון התראה "
             "כנקראה ולמחיקת התראה.")
    add_heading(doc, "תשלומים ודוחות", 2)
    add_para(doc,
             "פרטי התשלום נשמרים בטבלת Payments. המורה יכול לראות תשלומים לפי תלמיד, לאשר תשלום "
             "ולהפיק דוח לתקופה. התלמיד יכול לראות תשלומים פתוחים ולבדוק מה נשאר לשלם.")
    add_heading(doc, "בדיקות קלט ואבטחה", 2, "security")
    add_para(doc,
             "בזמן הרשמה והתחברות אני לא שומר סיסמה רגילה במסד. הסיסמה עוברת גיבוב עם PBKDF2 ו-salt. "
             "בכניסה למערכת משווים בין הערך המחושב לבין הערך השמור.")
    add_para(doc,
             "בנוסף, השאילתות משתמשות ב-OleDbParameter. כך הנתונים של המשתמש נשלחים כפרמטרים ולא מתחברים "
             "ישירות לטקסט של השאילתה. זה מקטין את הסיכון ל-SQL Injection.")
    add_code(doc, """
string sql = "SELECT * FROM [Lessons] WHERE StudentID = ?";
return Select(sql, new OleDbParameter("@sid", sid));

string hashedPassword = SecurityHelper.HashPassword(user.Password);""",
             "קטע קוד: פרמטר בשאילתה וגיבוב סיסמה לפני שמירה.")

    add_page_break(doc)
    add_heading(doc, "פרק 6: מדריך הפעלה ובדיקות", 1, "run")
    add_heading(doc, "הפעלה קצרה", 2)
    add_number(doc, "פותחים את הפתרון WcfServiceLibrary1.sln בתוכנת Visual Studio.")
    add_number(doc, "מוודאים שקובץ UsersDataBase.accdb נמצא בתוך תיקיית ViewDB.")
    add_number(doc, "מריצים את שירות WCF.")
    add_number(doc, "מריצים את הלקוח הרצוי: WPF, Web או MAUI.")
    add_number(doc, "נכנסים עם משתמש מתאים ובודקים את הפעולה הרצויה.")
    add_heading(doc, "בדיקות שביצעתי", 2, "tests")
    add_table(doc, ["בדיקה", "מה בדקתי", "תוצאה צפויה"], [
        ("התחברות", "שם משתמש וסיסמה תקינים", "מעבר למסך שמתאים לתפקיד."),
        ("בדיקת הרשאה", "תלמיד ומורה נכנסים למסכים שונים", "כל משתמש רואה פעולות מתאימות."),
        ("קביעת שיעור", "תלמיד בוחר תאריך ושעה", "נשמר שיעור חדש במסד."),
        ("ביטול שיעור", "ביטול שיעור קיים", "מצב השיעור מתעדכן."),
        ("תשלום", "אישור תשלום על שיעור", "מצב התשלום מתעדכן."),
        ("התראות", "שליפה וסימון כנקרא", "מוצגות רק התראות של המשתמש."),
        ("פניות תמיכה", "פתיחת פנייה והוספת הודעה", "הפנייה נשמרת עם ההודעות."),
    ], [1.35, 2.55, 2.30], font_size=9.2)
    add_heading(doc, "הערה לגבי מסכי הקורסים", 2)
    add_para(doc,
             "בפרויקט קיימות טבלאות ומחלקות עבור קורסים, אבל חלק ממסכי הקורסים עדיין לא פתוחים למשתמש. "
             "ביישום WPF סימנתי אותם כ-Coming Soon כדי לא להציג מסך חלקי כאילו הוא מוכן.")

    add_page_break(doc)
    add_heading(doc, "פרק 7: נספחים", 1, "appendix")
    add_para(doc,
             "בגוף התיק הכנסתי רק קטעי קוד קצרים שעוזרים להבין את המימוש. קוד המקור המלא מצורף "
             "בסוף קובץ ההגשה המלא, אחרי העמוד האחרון של התיק.")
    add_heading(doc, "קבצים שכדאי לפתוח בזמן הבחינה", 2)
    add_bullet(doc, "WcfServiceLibrary1/WcfServiceLibrary1/IService1.cs")
    add_bullet(doc, "WcfServiceLibrary1/WcfServiceLibrary1/Service1.cs")
    add_bullet(doc, "WcfServiceLibrary1/WcfServiceLibrary1/App.config")
    add_bullet(doc, "WcfServiceLibrary1/ViewDB/BaseDB.cs")
    add_bullet(doc, "WcfServiceLibrary1/ViewDB/LessonsDB.cs")
    add_bullet(doc, "WcfServiceLibrary1/Model/Helpers/SecurityHelper.cs")
    add_bullet(doc, "driver-client/driver-client/StudentUI.xaml")
    add_bullet(doc, "driver-maui/Pages/LoginPage.xaml.cs")
    add_heading(doc, "סיכום קצר", 2)
    add_para(doc,
             "Driver Moodle מרכז את ניהול שיעורי הנהיגה במקום אחד. בניתי מערכת לקוח-שרת עם בסיס נתונים, "
             "שירות WCF ושלושה לקוחות. במהלך העבודה למדתי לחלק פרויקט לשכבות, לעבוד עם מסד נתונים "
             "ולהתאים ממשק משתמש לסוגים שונים של משתמשים.")
    doc.save(OUT)
    print(OUT)


def build_qa():
    doc = Document()
    configure_doc(doc)
    add_para(doc, "Driver Moodle", size=22, bold=True, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, rtl=False, after=3)
    add_para(doc, "שאלות קשות שהבוחן יכול לשאול", size=17, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_para(doc,
             "הדף הזה מיועד להתכוננות בעל פה. התשובות קצרות בכוונה, כדי שיהיה קל לזכור אותן.",
             color=GRAY, after=10)
    qa = [
        ("למה בחרת לבנות את המערכת?",
         "בזמן שלמדתי נהיגה היה לי קשה לעקוב אחרי קביעת שיעורים, שינויים ותשלומים. רציתי לרכז את הכול במקום אחד."),
        ("למה זו מערכת לקוח-שרת?",
         "יש כמה סוגי לקוחות, אבל הנתונים צריכים להישאר במקום אחד. הלקוחות שולחים בקשות לשרת ורק השרת ניגש למסד."),
        ("למה השתמשת ב-WPF?",
         "WPF מתאים לעבודה מלאה מהמחשב הביתי. יש בו מקום למסכים גדולים ולפעולות רבות של תלמיד, מורה ומנהל."),
        ("למה הוספת Web?",
         "Web מאפשר להיכנס מדפדפן בלי התקנה. זה שימושי כשאין גישה למחשב שבו מותקן היישום."),
        ("למה הוספת MAUI / Android?",
         "מהטלפון נוח לבצע פעולות קצרות כמו בדיקת שיעורים, תשלום והתראות."),
        ("איפה נמצאות הגדרות הרשת?",
         "בקובצי App.config. שם נמצאת כתובת שירות WCF וסוג החיבור basicHttpBinding."),
        ("מה ההבדל בין IService1 לבין Service1?",
         "IService1 מגדיר אילו פעולות מותר לבקש מהשירות. Service1 מממש את הפעולות בפועל."),
        ("מה התפקיד של Model?",
         "Model מכיל מחלקות שמייצגות את הנתונים שעוברים בין הלקוחות לשרת, לדוגמה UserInfo ו-Notification."),
        ("מה התפקיד של ViewDB?",
         "ViewDB מרכז את העבודה מול Access. BaseDB מטפלת בחיבור ובמנגנוני שליפה ושמירה."),
        ("מהו מנגנון Select?",
         "זו פעולה משותפת ב-BaseDB שמקבלת שאילתה ופרמטרים, פותחת חיבור, קוראת את הרשומות ומחזירה אובייקטים."),
        ("איך הגנת מפני SQL Injection?",
         "השתמשתי ב-OleDbParameter. הערכים של המשתמש נשלחים כפרמטרים ולא מתחברים לטקסט של השאילתה."),
        ("איך הסיסמאות נשמרות?",
         "הן לא נשמרות כטקסט רגיל. אני משתמש ב-PBKDF2 עם salt ושומר את תוצאת הגיבוב."),
        ("מהו Async ולמה השתמשת בו?",
         "קריאה אסינכרונית מאפשרת להמתין לשרת בלי לתקוע את המסך. השתמשתי בזה בעיקר בלקוח MAUI."),
        ("מה זה IValueConverter?",
         "זו מחלקה שממירה מידע לצורה שמתאימה להצגה ב-WPF, בלי לערבב את נתוני המערכת עם עיצוב המסך."),
        ("למה יש קשרים רציפים ומקווקווים בתרשים המסד?",
         "קו רציף הוא קשר שמוגדר במסד. קו מקווקו הוא קישור לפי מזהים שמנוהל בקוד השירות."),
        ("האם כל מסכי הקורסים מוכנים?",
         "לא. התשתית קיימת, אבל חלק מהמסכים עדיין לא הושלמו. לכן סימנתי אותם כ-Coming Soon."),
        ("למה בחרת ב-Access ולא במסד גדול יותר?",
         "לפרויקט לימודי Access מספיק, קל להציג את הטבלאות והקשרים בזמן הבחינה, ולא צריך להתקין שרת מסד נפרד. במערכת אמיתית עם הרבה משתמשים הייתי עובר למסד שרת."),
        ("למה בנית שלושה לקוחות ולא לקוח אחד?",
         "כל הלקוחות משתמשים באותם נתונים ובאותו שירות, אבל כל אחד מתאים למצב אחר: WPF לעבודה מלאה במחשב, Web לכניסה מהירה מדפדפן, ו-MAUI לפעולות קצרות בטלפון."),
        ("מה קורה כאשר תלמיד קובע שיעור?",
         "הלקוח שולח בקשה לשירות WCF. Service1 מעביר את הפעולה ל-LessonsDB, ושכבת ViewDB שומרת את השיעור במסד הנתונים."),
        ("מה יקרה אם כתובת השרת תשתנה?",
         "לא צריך לשנות כל מסך. מעדכנים את הכתובת בקובץ App.config של הלקוח ובקובץ ההגדרות של השירות."),
        ("למה לא הגדרת מפתח זר לכל קו בתרשים?",
         "חלק מהקישורים נשמרים לפי מזהה וסוג משתמש, למשל בהתראות. כרגע השירות בודק אותם בקוד. לכן סימנתי אותם בקו מקווקו ולא הצגתי אותם כמפתח זר במסד."),
        ("מה ההבדל בין Validation לבין ניקוי קלט?",
         "Validation בודק אם הקלט תקין, למשל אם כתובת דואר אלקטרוני כתובה בצורה נכונה. ניקוי קלט מסיר תווים לא רצויים לפני שמירה או הצגה."),
        ("איך נשמרת ההפרדה בין סוגי המשתמשים?",
         "לאחר הכניסה נשמר סוג המשתמש, ולפי התפקיד מוצגים המסכים המתאימים. בפעולות רגישות חשוב לבדוק הרשאה גם בשירות ולא להסתמך רק על המסך."),
        ("למה קוד המקור המלא נמצא בסוף ולא בתוך כל פרק?",
         "בגוף התיק שמתי רק קטעים קצרים שעוזרים להבין את הרעיון. בסוף מצורף הקוד המלא כדי שהבוחן יוכל לפתוח כל קובץ בלי להעמיס על ההסבר."),
        ("מה היית משפר אם היית ממשיך את הפרויקט?",
         "הייתי משלים את מסכי הקורסים, מוסיף עוד בדיקות הרשאה בצד השירות, ומעביר את מסד הנתונים למסד שרת אם מספר המשתמשים היה גדל."),
    ]
    for idx, (q, a) in enumerate(qa, 1):
        add_heading(doc, f"{idx}. {q}", 2)
        add_para(doc, a, after=7)
    doc.save(QA_OUT)
    print(QA_OUT)


if __name__ == "__main__":
    build_main()
    build_qa()
